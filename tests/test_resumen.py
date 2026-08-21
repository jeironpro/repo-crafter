from io import BytesIO
from unittest.mock import patch

from docx import Document

import resumen
from tests.conftest import RespuestaFake


def test_topics_despliegue_detecta_cualquier_topic_con_pages(repo_falso):
    repo = repo_falso(topics=["flask", "github-pages", "cloudflare-pages", "api"])
    assert resumen.topics_despliegue(repo) == ["github-pages", "cloudflare-pages"]


def test_esta_desplegado_es_insensible_a_mayusculas(repo_falso):
    assert resumen.esta_desplegado(repo_falso(topics=["GitHub-Pages"]))
    assert not resumen.esta_desplegado(repo_falso(topics=["frontend"]))


def test_clasificar_repos_separa_en_tres_grupos(repo_falso):
    repos = [
        repo_falso(nombre="desplegado-publico", privado=False, topics=["github-pages"]),
        repo_falso(nombre="desplegado-privado", privado=True, topics=["cloudflare-pages"]),
        repo_falso(nombre="normal-publico", privado=False),
        repo_falso(nombre="normal-privado", privado=True),
    ]

    grupos = resumen.clasificar_repos(repos)

    nombres = {grupo: [r["name"] for r in lista] for grupo, lista in grupos.items()}
    assert set(nombres["desplegados"]) == {"desplegado-publico", "desplegado-privado"}
    assert nombres["publicos"] == ["normal-publico"]
    assert nombres["privados"] == ["normal-privado"]


def test_clasificar_repos_ordena_alfabeticamente(repo_falso):
    repos = [
        repo_falso(nombre="zeta"),
        repo_falso(nombre="Alfa"),
        repo_falso(nombre="media"),
    ]

    grupos = resumen.clasificar_repos(repos)

    assert [r["name"] for r in grupos["publicos"]] == ["Alfa", "media", "zeta"]


def test_nombre_archivo_incluye_fecha():
    assert resumen.nombre_archivo().startswith("resumen-repos-")
    assert resumen.nombre_archivo().endswith(".pdf")
    assert resumen.nombre_archivo("docx").endswith(".docx")


def test_generar_docx_produce_documento_con_tres_tablas(repo_falso):
    repos = [
        repo_falso(nombre="b-pages", topics=["github-pages"]),
        repo_falso(nombre="a-privado", privado=True),
    ]
    grupos = resumen.clasificar_repos(repos)
    topics = {r["name"]: resumen.topics_despliegue(r) for r in repos}

    documento = Document(BytesIO(resumen.generar_docx(grupos, topics, "usuario-test")))

    assert len(documento.tables) == 3

    desplegados = documento.tables[0]
    assert [celda.text for celda in desplegados.rows[0].cells] == [
        "Nombre", "Visibilidad", "Topics de despliegue"
    ]
    assert desplegados.rows[1].cells[0].text == "b-pages"
    assert desplegados.rows[1].cells[1].text == "Público"
    assert "github-pages" in desplegados.rows[1].cells[2].text

    assert documento.tables[1].rows[1].cells[0].text == "Ninguno"
    assert documento.tables[2].rows[1].cells[0].text == "a-privado"


def test_ruta_resumen_formato_docx_devuelve_word_valido(cliente, repo_falso):
    repos = [repo_falso(nombre="con-pages", topics=["github-pages"])]

    with patch("github_api.requests.get", side_effect=[RespuestaFake(datos=repos), RespuestaFake(datos=[])]):
        respuesta = cliente.get("/resumen?formato=docx")

    assert respuesta.status_code == 200
    assert "wordprocessingml" in respuesta.mimetype
    assert resumen.nombre_archivo("docx") in respuesta.headers["Content-Disposition"]

    documento = Document(BytesIO(respuesta.data))
    assert len(documento.tables) == 3
    assert documento.tables[0].rows[1].cells[0].text == "con-pages"


def test_ruta_resumen_formato_invalido_rechaza(cliente):
    assert cliente.get("/resumen?formato=xls").status_code == 400


def test_ruta_resumen_devuelve_pdf(cliente, repo_falso):
    repos = [
        repo_falso(nombre="con-pages", topics=["github-pages"]),
        repo_falso(nombre="sin-pages"),
        repo_falso(nombre="privado-sin-pages", privado=True),
    ]

    with patch("github_api.requests.get", side_effect=[RespuestaFake(datos=repos), RespuestaFake(datos=[])]), \
         patch("resumen.generar_pdf", return_value=b"%PDF-fake") as generar:
        respuesta = cliente.get("/resumen")

    assert respuesta.status_code == 200
    assert respuesta.mimetype == "application/pdf"
    assert respuesta.data == b"%PDF-fake"
    assert "attachment" in respuesta.headers["Content-Disposition"]
    assert resumen.nombre_archivo() in respuesta.headers["Content-Disposition"]

    html_renderizado = generar.call_args[0][0]
    assert "Repositorios desplegados" in html_renderizado
    assert "Repositorios públicos no desplegados" in html_renderizado
    assert "Repositorios privados no desplegados" in html_renderizado
    assert "con-pages" in html_renderizado


def test_ruta_resumen_sin_repos_devuelve_pdf_vacio_de_datos(cliente):
    with patch("github_api.requests.get", return_value=RespuestaFake(datos=[])), \
         patch("resumen.generar_pdf", return_value=b"%PDF-vacio") as generar:
        respuesta = cliente.get("/resumen")

    assert respuesta.status_code == 200
    html_renderizado = generar.call_args[0][0]
    assert html_renderizado.count("Ninguno") == 3

"""Clasificación de repositorios y generación del resumen en PDF y DOCX.

Un repo se considera desplegado cuando algún topic contiene "pages"
(p. ej. github-pages, cloudflare-pages, cf-pages).
"""
import io
from datetime import date

from docx import Document
from weasyprint import HTML

TOPIC_DESPLIEGUE = "pages"

SECCIONES = [
    ("Repositorios desplegados", "desplegados", True),
    ("Repositorios públicos no desplegados", "publicos", False),
    ("Repositorios privados no desplegados", "privados", False),
]


def topics_despliegue(repo):
    return [topic for topic in repo.get("topics", []) if TOPIC_DESPLIEGUE in topic.lower()]


def esta_desplegado(repo):
    return bool(topics_despliegue(repo))


def clasificar_repos(repos):
    grupos = {"desplegados": [], "publicos": [], "privados": []}

    for repo in repos:
        if esta_desplegado(repo):
            grupos["desplegados"].append(repo)
        elif repo["private"]:
            grupos["privados"].append(repo)
        else:
            grupos["publicos"].append(repo)

    for lista in grupos.values():
        lista.sort(key=lambda repo: repo["name"].lower())

    return grupos


def nombre_archivo(formato="pdf"):
    return f"resumen-repos-{date.today().isoformat()}.{formato}"


def generar_pdf(html):
    return HTML(string=html).write_pdf()


def generar_docx(grupos, topics_despliegue, usuario):
    """Construye el resumen como documento Word editable con el mismo contenido que el PDF."""
    documento = Document()
    documento.add_heading("Resumen de repositorios", 0)

    total = sum(len(lista) for lista in grupos.values())
    documento.add_paragraph(
        f"@{usuario} · {date.today().strftime('%d/%m/%Y')} · {total} repositorios"
    )

    for titulo, clave, con_topics in SECCIONES:
        repos = grupos[clave]
        documento.add_heading(titulo, level=1)

        columnas = ["Nombre", "Visibilidad", "Topics de despliegue"] if con_topics else ["Nombre"]
        tabla = documento.add_table(rows=1, cols=len(columnas))
        tabla.style = "Light Grid Accent 1"

        for indice, encabezado in enumerate(columnas):
            tabla.rows[0].cells[indice].text = encabezado

        if repos:
            for repo in repos:
                celdas = tabla.add_row().cells
                celdas[0].text = repo["name"]
                if con_topics:
                    celdas[1].text = "Privado" if repo["private"] else "Público"
                    celdas[2].text = ", ".join(topics_despliegue.get(repo["name"], []))
        else:
            tabla.add_row().cells[0].text = "Ninguno"

    buffer = io.BytesIO()
    documento.save(buffer)
    return buffer.getvalue()
